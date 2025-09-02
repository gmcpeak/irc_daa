# app.py
import os
import io
import base64
from pathlib import Path
from typing import Optional, List
from datetime import datetime

import streamlit as st
from PIL import Image, ImageOps
from PIL.Image import Image as PILImage
from openai import OpenAI
from docx import Document  # NEW: for .docx output

# -----------------------------
# Config & constants
# -----------------------------
st.set_page_config(page_title="GW+IRC DAA", layout="wide")

PROMPT_PATH = "./assets/system_prompt.txt"
DOC1_PATH   = "./assets/Goodwill-Donation-Value-Guide.txt"
DOC2_PATH   = "./assets/Salvation-Army-Donation-Value-Guide.txt"

openai_key = st.secrets.get("openai_key")  # safer get()
client = OpenAI(api_key=openai_key) if openai_key else None

# -----------------------------
# Session state (for download)
# -----------------------------
if "last_output" not in st.session_state:
    st.session_state.last_output = None
if "docx_bytes" not in st.session_state:
    st.session_state.docx_bytes = None
if "docx_filename" not in st.session_state:
    st.session_state.docx_filename = None

# -----------------------------
# Helpers
# -----------------------------
def _read_texts(paths: List[str]) -> str:
    """Read all existing text files from paths and concatenate. Silently skip missing."""
    parts: List[str] = []
    for p in paths:
        try:
            with open(p, "r", encoding="utf-8", errors="ignore") as f:
                parts.append(f.read())
        except Exception:
            pass
    return "\n".join(parts)

def _to_mm_content(text: Optional[str], image: Optional[PILImage | str]):
    """
    Build multimodal content list for the Responses API.
    Accepts text and either a PIL Image or a data-URL string.
    """
    content = []
    if text:
        content.append({"type": "input_text", "text": text})

    if isinstance(image, str) and image.startswith("data:image"):
        content.append({"type": "input_image", "image_url": image})
    elif isinstance(image, PILImage):
        try:
            buf = io.BytesIO()
            image.save(buf, format="PNG")
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            content.append({"type": "input_image", "image_url": f"data:image/png;base64,{b64}"})
        except Exception:
            pass

    return content

def _load_pil_from_uploaded(uploaded) -> Optional[PILImage]:
    """st.camera_input / st.file_uploader both return UploadedFile (file-like)."""
    if uploaded is None:
        return None
    try:
        img = Image.open(uploaded)
        return ImageOps.exif_transpose(img)
    except Exception:
        try:
            data = uploaded.getvalue()
            return ImageOps.exif_transpose(Image.open(io.BytesIO(data)))
        except Exception:
            return None

def _process(prompt_text: str, image: Optional[PILImage]) -> str:
    """Call OpenAI Responses API with multimodal input and web_search tool."""
    if client is None:
        return "OpenAI client not configured (missing API key)."

    content = _to_mm_content(prompt_text, image)
    response = client.responses.create(
        model="gpt-5",
        reasoning={"effort": "low"},
        text={"verbosity": "low"},
        tools=[{"type": "web_search"}],
        input=[{"role": "user", "content": content}] if content else prompt_text,
    )

    text = getattr(response, "output_text", None)
    if not text:
        try:
            text = response.output[0].content[0].text
        except Exception:
            text = str(response)
    return text

def _docx_bytes_from_text(
    text: str,
    image: Optional[PILImage],
    title: str = "Donation Audit",
) -> bytes:
    """
    Build a .docx in memory. Adds a heading, then the original image
    (scaled to fit within page margins), then the generated text.
    """
    doc = Document()

    # Heading
    if title:
        doc.add_heading(title, level=0)

    # Image (after header, before text)
    if image is not None:
        img_buf = io.BytesIO()
        image.save(img_buf, format="PNG")
        img_buf.seek(0)  # important when using BytesIO
        section = doc.sections[0]
        max_width = section.page_width - section.left_margin - section.right_margin
        doc.add_picture(img_buf, width=max_width)
        doc.add_paragraph("")  # a little space after the image

    # Text (preserve blank lines)
    for line in (text or "").splitlines():
        doc.add_paragraph(line)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# -----------------------------
# UI
# -----------------------------
st.markdown("# GW+IRC Donation Audit Assistant")
st.markdown("Use your phone or laptop camera to take a picture **or** upload one, then click **Process**.")

left, right = st.columns(2)
with left:
    cam = st.camera_input(
        "Take a photo",
        label_visibility="visible",
        key="camera",
        help="Use your device camera",
        disabled=False
    ) 
with right:
    upl = st.file_uploader(
        "...or upload a photo",
        type=["png", "jpg", "jpeg"],
        accept_multiple_files=False,
        key="uploader"
    )

uploaded = cam or upl
image_pil = _load_pil_from_uploaded(uploaded)

st.markdown("### Output")
out_area = st.empty()

process_clicked = st.button(
    "Process",
    type="primary",
    disabled=(image_pil is None),  # only selectable when an image is present
    help="Upload or take a photo first"
)

if process_clicked:
    system_prompt = _read_texts([PROMPT_PATH])
    guides_text   = _read_texts([DOC1_PATH, DOC2_PATH])
    prompt = (system_prompt or "") + "\n" + (guides_text or "")

    with st.spinner("Generating output..."):
        try:
            out_text = _process(prompt, image_pil)
        except Exception as e:
            out_text = f"Error calling OpenAI: {e}"

    # Show output
    out_area.markdown(out_text if out_text else "_No output._")

    # Build .docx bytes and stash in session_state for the next rerun
    st.session_state.last_output = out_text
    st.session_state.docx_bytes = _docx_bytes_from_text(out_text or "", image_pil)
    st.session_state.docx_filename = f"GW-IRC-Donation-Value-Audit--{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"

# Show download button only after we have generated output
if st.session_state.docx_bytes and st.session_state.last_output:
    st.download_button(
        label="Download as Word (.docx)",
        data=st.session_state.docx_bytes,
        file_name=st.session_state.docx_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        help="Save the generated output as a Word document",
        key="download_docx"
    )
